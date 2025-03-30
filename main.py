import os, asyncio, jwt, bcrypt, cloudinary.uploader, pandas as pd
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, Form, BackgroundTasks, Depends, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from docx import Document
from docx.shared import Pt, RGBColor
from spire.doc import Document as SpireDocument, FileFormat
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
CLOUD_NAME, API_KEY, API_SECRET, SECRET_KEY = map(os.getenv, ["CLOUD_NAME", "API_KEY", "API_SECRET", "SECRET_KEY"])
ALLOWED_USERS, revoked_tokens, progress_data = eval(os.getenv("ALLOWED_USERS", "{}")), set(), {}

# Configure Cloudinary
cloudinary.config(cloud_name=CLOUD_NAME, api_key=API_KEY, api_secret=API_SECRET)

# FastAPI setup
app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

FONT_STYLES = {"{name}": Pt(26), "{department}": Pt(21), "{year}": Pt(21), "{event}": Pt(20.5), "{date}": Pt(19)}

# Utility functions
def upload_to_cloudinary(file_path, folder):
    result = cloudinary.uploader.upload(file_path, folder=folder, use_filename=True, unique_filename=False, resource_type="raw" if file_path.endswith("_logs.xlsx") else "auto")
    os.remove(file_path)
    return result["secure_url"]

def convert_docx_to_pdf(docx_path, event):
    pdf_path = docx_path.with_suffix(".pdf")
    doc = SpireDocument()
    doc.LoadFromFile(str(docx_path))
    doc.SaveToFile(str(pdf_path), FileFormat.PDF)
    os.remove(docx_path)
    return upload_to_cloudinary(str(pdf_path), f"AutoCred/{event}")

def generate_certificate(template, placeholders, output_folder, event, role):
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    doc, name, email = Document(template), placeholders["{name}"].replace(" ", "_"), placeholders["{email}"].replace(" ", "_")
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, "")
                            run = paragraph.add_run(value)
                            run.font.size, run.font.italic, run.font.color.rgb = FONT_STYLES.get(key, Pt(21.5)), True, RGBColor(171, 124, 52)

    docx_path = Path(output_folder) / f"{name}_{email}_{role}.docx"
    doc.save(docx_path)
    return convert_docx_to_pdf(docx_path, event)

async def process_bulk_certificates(event_name, event_date, template_path, file_path, pdf_folder, role):
    data = pd.read_excel(file_path, engine="openpyxl").to_dict(orient="records")
    log_file_path, results = Path(file_path).with_name(f"{event_name}_logs.xlsx"), []
    progress_data[event_name] = {"completed": 0, "total": len(data)}
    
    for student in data:
        placeholders = {f"{{{k.lower()}}}": v for k, v in student.items()}
        placeholders.update({"{event}": event_name, "{date}": event_date})
        results.append({"Name": student["Name"], "Email": student["Email"], "Certificate": generate_certificate(template_path, placeholders, pdf_folder, event_name, role)})
        progress_data[event_name]["completed"] += 1
        await asyncio.sleep(0.1)

    pd.DataFrame(results).to_excel(log_file_path, index=False)
    del progress_data[event_name]
    return upload_to_cloudinary(str(log_file_path), f"AutoCred/{event_name}")

# Authentication functions
def verify_password(plain_password, hashed_password):
    return bcrypt.checkpw(plain_password.encode(), hashed_password.encode())

def create_access_token(email):
    return jwt.encode({"sub": email}, SECRET_KEY, algorithm="HS256")

async def get_current_user(token: str = Depends(oauth2_scheme)):
    try:
        if token in revoked_tokens:
            raise HTTPException(status_code=401, detail="Token revoked")
        email = jwt.decode(token, SECRET_KEY, algorithms=["HS256"]).get("sub")
        if email not in ALLOWED_USERS:
            raise HTTPException(status_code=401, detail="Invalid token")
        return email
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Could not validate credentials")

# Routes
@app.post("/token")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    email, password = form_data.username, form_data.password
    if email not in ALLOWED_USERS or not verify_password(password, ALLOWED_USERS[email]):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Incorrect email or password")
    return {"access_token": create_access_token(email), "token_type": "bearer"}

@app.post("/logout")
async def logout(token: str = Depends(oauth2_scheme)):
    revoked_tokens.add(token)
    return {"message": "Successfully logged out"}

@app.get("/progress/{event_name}/total")
async def get_total_certificates(event_name: str):
    return {"total_certificates": progress_data.get(event_name, {}).get("total", "Event not found")}

@app.get("/progress/{event_name}/completed")
async def get_completed_certificates(event_name: str):
    if event_name not in progress_data:
        return JSONResponse({"error": "Event not found"}, status_code=404)

    async def event_stream():
        while event_name in progress_data:
            yield f"data: {progress_data[event_name]['completed']}\n\n"
            await asyncio.sleep(1)

    return StreamingResponse(event_stream(), media_type="text/event-stream")

@app.post("/generate-certificates")
async def generate_certificates(
    background_tasks: BackgroundTasks,
    event_name: str = Form(...),
    event_date: str = Form(...),
    template: str = Form(...),
    gen_type: str = Form(...),
    file: UploadFile = File(None),
    student_name: str = Form(None),
    department: str = Form(None),
    year: str = Form(None),
    email: str = Form(None),
    current_user: str = Depends(get_current_user),
):
    script_dir, role = Path(__file__).parent, "participant" if template == "template1" else "organizer"
    template_path, pdf_folder = script_dir / f"temp{1 if template == 'template1' else 2}.docx", script_dir / "tmp/Certificates"
    pdf_folder.mkdir(parents=True, exist_ok=True)

    if gen_type == "single":
        placeholders = {k: v for k, v in zip(["{name}", "{department}", "{year}", "{event}", "{date}", "{email}"], [student_name, department, year, event_name, event_date, email])}
        return {"message": "Certificate generated successfully", "download_url": generate_certificate(template_path, placeholders, pdf_folder, event_name, role)}

    if gen_type == "bulk" and file:
        excel_path = script_dir / file.filename
        with excel_path.open("wb") as f:
            f.write(file.file.read())

        log_file_url = await process_bulk_certificates(event_name, event_date, template_path, excel_path, pdf_folder, role)
        return {"message": "Bulk certificate generation completed", "log_file_url": log_file_url}

    return JSONResponse({"error": "Invalid generation type or missing data"}, status_code=400)
