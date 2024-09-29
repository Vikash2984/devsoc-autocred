import csv
from docx import Document
from docx.shared import Pt, RGBColor
import os

def read_student_data(csv_file):
    """
    Reads student data (name, department, and year) from a CSV file and returns a list of dictionaries.
    """
    student_data = []
    if os.path.isfile(csv_file):  # Check if file exists
        with open(csv_file, mode='r', newline='', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            required_columns = ['Name', 'Department', 'Year']
            # Ensure required columns are present in the CSV file
            if all(column in reader.fieldnames for column in required_columns):
                for row in reader:
                    # Append each row as a dictionary to the list
                    student_data.append({
                        'Name': row['Name'],
                        'Department': row['Department'],
                        'Year': row['Year']
                    })
            else:
                print(f"CSV file must contain the columns: {', '.join(required_columns)}")
    else:
        print(f"File not found: {csv_file}")
    return student_data

def replace_placeholder_with_format(paragraph, placeholder, replacement):
    """
    Replaces a placeholder in the paragraph with the replacement text formatted as specified.
    """
    for run in paragraph.runs:
        if placeholder in run.text:
            # Create a new run with the desired formatting
            run.text = run.text.replace(placeholder, "")  # Remove placeholder text from this run
            new_run = paragraph.add_run(replacement)  # Add new run with the replacement text
            new_run.font.name = 'Georgia'  # Set the font to Georgia
            new_run.font.size = Pt(21.5)  # Set the font size
            new_run.font.color.rgb = RGBColor(171, 124, 52)  # Set the font color (#AB7C34 in RGB)
            new_run.font.italic = True  # Make the font italic

def replace_placeholders_in_table(table, placeholders):
    """
    Replace placeholders in all cells of a table with the specified replacements.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for placeholder, replacement in placeholders.items():
                    replace_placeholder_with_format(paragraph, placeholder, replacement)

def generate_certificate(template_path, output_folder, student_data):
    """
    Generates certificates for students using the given Word template and student data.
    Replaces {name}, {department}, and {year} in the template.
    """
    for student in student_data:
        # Load the Word template
        doc = Document(template_path)

        # Define placeholders and their replacements
        placeholders = {
            '{name}': student['Name'],
            '{department}': student['Department'],
            '{year}': student['Year']
        }

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in placeholders.items():
                replace_placeholder_with_format(paragraph, placeholder, replacement)

        # Replace placeholders in tables
        for table in doc.tables:
            replace_placeholders_in_table(table, placeholders)

        # Save the modified document
        output_path = os.path.join(output_folder, f"{student['Name']}_certificate.docx")
        doc.save(output_path)
        print(f"Generated certificate for {student['Name']} at '{output_path}'")

# Paths
csv_file = r'Generator-word\CSV-Generator\data.csv'
template_path = r'template.docx'
output_folder = r'docx'

# Read student data from CSV
student_data = read_student_data(csv_file)

# Generate certificates
generate_certificate(template_path, output_folder, student_data)
