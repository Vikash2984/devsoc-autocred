from docx import Document
from docx.shared import Pt, RGBColor
import pandas as pd
import os

def read_student_data(excel_file):
    """
    Reads student data (name, department, and year) from an Excel file and returns a list of dictionaries.
    Converts all data to strings where necessary.
    """
    student_data = []
    if os.path.isfile(excel_file):  # Check if file exists
        df = pd.read_excel(excel_file, engine='openpyxl')
        # Ensure the required columns are present in the Excel file
        required_columns = ['Name', 'Email', 'Department', 'Year']
        if all(column in df.columns for column in required_columns):
            # Convert all data to string and return it
            student_data = df[required_columns].astype(str).to_dict(orient='records')
        else:
            print(f"Excel file must contain the columns: {', '.join(required_columns)}")
    else:
        print(f"File not found: {excel_file}")
    return student_data

def replace_placeholder_with_format(paragraph, placeholder, replacement):
    """
    Replaces a placeholder in the paragraph with the replacement text formatted as specified.
    """
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, "")  # Remove placeholder text from this run
            new_run = paragraph.add_run(replacement)  # Add new run with the replacement text
            new_run.font.name = 'Georgia'  # Set the font to Georgia
            new_run.font.color.rgb = RGBColor(171, 124, 52)  # Set the font color (#AB7C34 in RGB)
            new_run.font.italic = True  # Make the font italic
            
            # Conditionally set font size for specific placeholders
            if placeholder == '{name}':
                new_run.font.size = Pt(28)  # Set font size to 28 for 'name'
            else:
                new_run.font.size = Pt(21.5)  # Keep font size 21.5 for 'department' and 'year'

def replace_placeholders_in_table(table, placeholders):
    """
    Replace placeholders in all cells of a table with the specified replacements.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for placeholder, replacement in placeholders.items():
                    replace_placeholder_with_format(paragraph, placeholder, replacement)

def generate_certificate(template_path, output_folder, student_data):
    """
    Generate certificates by replacing placeholders with actual student data.
    """
    gen_count = 0
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    for student in student_data:
        placeholders = {
            '{name}': student['Name'],
            '{department}': student['Department'],
            '{year}': student['Year']
        }
        
        doc = Document(template_path)
        for table in doc.tables:
            replace_placeholders_in_table(table, placeholders)
        
        output_path = os.path.join(output_folder, f"{student['Name']}_{student['Email']}_certificate.docx")
        doc.save(output_path)
        print(f"\nCertificate generated for {student['Name']} at {output_path}")
        gen_count += 1
    if gen_count == 1:
        print(f"\n 1 certificate was generated...")
    else:
        print(f"\n{gen_count} certificates were generated...")

# Example usage:
template_path = "template.docx"
output_folder = "docx"
excel_file = "Generator-word\XLSX-Generator\data.xlsx"
student_data = read_student_data(excel_file)
generate_certificate(template_path, output_folder, student_data)
